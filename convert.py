#!/usr/bin/env python3
"""扫描版 PDF -> 保版式可编辑 .docx  (全本地, 无需联网/API)

流水线:  PyMuPDF 渲染 -> RapidOCR 识别(带坐标) -> 版面重建 -> python-docx

用法:
    python3 convert.py 文件.pdf [更多.pdf ...]     # 端到端
    python3 convert.py --rebuild                   # 仅用已有 OCR 缓存重建 docx(调版式时用)

依赖:  pip install pymupdf rapidocr onnxruntime python-docx
"""
import json, glob, os, re, sys
import argparse, collections, traceback
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsmap

CACHE = '.pdfrec_cache'          # 页图与 OCR 结果缓存(运行时按 PDF 所在目录重设)

# ---------- 可调参数 (GUI 参数面板直接映射到这里) ----------
DEFAULTS = {
    'long_edge':   2560,    # 渲染长边像素: 越大越准也越慢
    'header_y':    0.125,   # 此线以上视为页眉区
    'footer_y':    0.915,   # 此线以下视为页脚区
    'gutter':      0.035,   # 行内横向空白超过此比例(占页宽)才算跨列 -> 表格行
    'min_tbl_rows': 2,      # 表格区至少要有这么多"跨列行"才做成表格
    'line_tol':    0.45,    # 视觉行聚类: y 重叠比例阈值
    'x_tol':       0.025,   # 同列判定的 x 容差(占页宽)
    'full_line':   0.78,    # 行右端超过此比例才算"排满"(才可能有续行)
    'bullet_ind':  0.030,   # 缩进超出基准多少算项目符号
    'stamp_conf':  0.88,    # 低于此置信度的短串视为印章/手写, 剔除
    'drop_header': True,    # 剔除页眉
    'drop_footer': True,    # 剔除页脚
    'drop_stamp':  True,    # 剔除印章/签名噪声
    'page_marker': True,    # 插入"—— 原第 N 页 ——"分隔标记
    'tables':      True,    # 启用多列版面 -> 无边框表格还原
    'grid_tables': True,    # 画了框线的表格照框线还原(含合并单元格)
    'debug':       False,   # 生成版面调试图(框出识别结果与块划分)
    'zh_font':     '宋体',
    'en_font':     'Times New Roman',
    'font_size':   10.5,
}
CFG = dict(DEFAULTS)

def G(k):
    return CFG.get(k, DEFAULTS[k])

FROZEN = getattr(sys, 'frozen', False)     # PyInstaller 打包后为 True

_last_tag = ''     # 最近一次转换的文件标签(GUI 预览据此定位页图)

class Cancelled(Exception):
    """用户中断转换"""

def _tick(progress, stage, cur, total, msg=''):
    if progress:
        progress(stage, cur, total, msg)

def _check(stop):
    if stop and stop():
        raise Cancelled()

def _log(log, msg):
    """统一日志出口: GUI 传回调收进日志区, 命令行则直接打印"""
    if log:
        log(msg)
    else:
        print(msg, flush=True)

def _record(errors, log, stage, page, exc):
    """记录单页失败, 不中断整篇 (页级异常隔离)"""
    info = {'stage': stage, 'page': page,
            'error': f'{type(exc).__name__}: {exc}',
            'trace': traceback.format_exc()}
    if errors is not None:
        errors.append(info)
    _log(log, f'  ! {stage}失败 (第 {page} 页), 已跳过: {info["error"]}')
    return info

def res_dir(*parts):
    """资源目录: 打包后指向解包目录 sys._MEIPASS, 否则指向脚本所在目录"""
    base = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, *parts)

def model_root():
    """离线模型目录; 打包时随 exe 分发, 避免首次运行联网下载"""
    for cand in (res_dir('models'), res_dir('rapidocr', 'models')):
        if os.path.isdir(cand) and glob.glob(os.path.join(cand, '*.onnx')):
            return cand
    return None                  # 未随包分发 -> 交给 rapidocr 自行下载

# ---------- 0. 渲染 + OCR ----------
def render_pdf(pdf, tag, progress=None, stop=None, errors=None, log=None):
    """按页面实际旋转渲染为 PNG; 自动选 dpi 使长边 ~G('long_edge')

    单页渲染失败(损坏的图像流等)只跳过该页, 不影响整篇。
    """
    import pymupdf
    d = pymupdf.open(pdf)
    outs = []
    n = d.page_count
    for i, p in enumerate(d):
        _check(stop)
        _tick(progress, 'render', i + 1, n, f'渲染第 {i+1}/{n} 页')
        out = f'{CACHE}/pages/{tag}_{i+1:03d}.png'
        if os.path.exists(out):
            outs.append(out)
            continue
        try:
            long_pt = max(p.rect.width, p.rect.height)
            dpi = max(150, min(300, round(G('long_edge') / long_pt * 72)))
            p.get_pixmap(dpi=dpi).save(out)     # get_pixmap 已应用 page.rotation
            outs.append(out)
        except Cancelled:
            raise
        except Exception as e:
            _record(errors, log, '渲染', i + 1, e)
    d.close()
    _log(log, f'  渲染完成: {len(outs)}/{n} 页')
    return outs

_ENGINE = None

def get_engine():
    """OCR 引擎全局复用: 加载模型约 1-2 秒, 每次新建会拖慢批量转换"""
    global _ENGINE
    if _ENGINE is None:
        from rapidocr import RapidOCR
        root = model_root()
        _ENGINE = RapidOCR(params={'Global.model_root_dir': root} if root else None)
    return _ENGINE

def _ocr_result(engine, png):
    """跑一页 OCR, 归一化成 items 列表"""
    r = engine(png)
    items = []
    if r is not None and r.txts:
        for txt, box, sc in zip(r.txts, r.boxes, r.scores):
            xs = [float(q[0]) for q in box]; ys = [float(q[1]) for q in box]
            items.append({'t': txt, 'x0': min(xs), 'y0': min(ys),
                          'x1': max(xs), 'y1': max(ys), 's': float(sc)})
    items.sort(key=lambda d: (d['y0'], d['x0']))
    return items

def run_ocr(pngs, tag, progress=None, stop=None, errors=None, log=None):
    """逐页 OCR, 结果落 JSON 缓存

    只跑串行: onnxruntime 本身就把所有核心吃满, 再拆多进程只是互相抢 CPU
    (实测 11 页 13.7s -> 13.7s), 要限死每进程线程数才换得来 1.4x, 代价是
    spawn/pickle/freeze_support 一整套额外失败面 —— 不值。
    """
    todo = [f for f in pngs
            if not os.path.exists(f'{CACHE}/ocr/{os.path.basename(f)[:-4]}.json')]
    if not todo:
        _log(log, f'  OCR: {len(pngs)} 页全部命中缓存, 跳过识别')
        return sorted(glob.glob(f'{CACHE}/ocr/{tag}_*.json'))

    _tick(progress, 'ocr', 0, len(todo), '加载识别模型...')
    _log(log, f'  OCR: {len(todo)} 页待识别')
    engine = get_engine()
    for n, f in enumerate(todo, 1):
        _check(stop)
        stem = os.path.basename(f)[:-4]
        _tick(progress, 'ocr', n, len(todo), f'识别 {stem} ({n}/{len(todo)})')
        try:
            items = _ocr_result(engine, f)
        except Cancelled:
            raise
        except Exception as e:
            _record(errors, log, 'OCR', n, e)     # 单页识别失败不影响其余页
            continue
        json.dump({'page': stem, 'items': items},
                  open(f'{CACHE}/ocr/{stem}.json', 'w'), ensure_ascii=False)
        _log(log, f'  OCR [{n}/{len(todo)}] {stem}  {len(items)} 行')
    return sorted(glob.glob(f'{CACHE}/ocr/{tag}_*.json'))

HEADER_PAT = re.compile(
    r'^(TROESTER|EXCELLENCE IN EXTRUSION|Contract No\.|January \d+|Attachment \d|'
    r'De?a?tailed Technical Description|Page\s*\d+\s*/\s*\d+|正直诚信|山东泰开电缆有限公司$)')
FOOTER_PAT = re.compile(r'^(\d{1,3}\s*/\s*\d{1,3}$|\d{1,3}$|邮箱|地址|邮编|网址|http|tkdlsc)')

NUM_START = re.compile(
    r'^(\d{1,2}\.\d{1,2}(\.\d{1,2})?[\s、．.]|'      # 1.1  6.1.2
    r'\d{1,2}[\.、]\s*|'                             # 1.  2、 3、设备(无空格)
    r'[一二三四五六七八九十]+[、．.]|'                 # 一、
    r'[（(]\d{1,2}[）)]|'                            # (1)
    r'[A-Z][\.、]\s|'                                # A.
    r'[•·●○\*]\s*|'                                  # 项目符号
    r'[-–—]\s+)')                                    # 短横线子项
END_PUNCT = re.compile(r'[。．.：:；;!?！？）)】\]"”]$')
BULLET    = re.compile(r'^[•·●○\*]\s*|^[-–—]\s+')

def is_zh(s):
    return sum(1 for c in s if '一' <= c <= '鿿') >= 2

def clean(t):
    t = t.replace('　', ' ').strip()
    t = re.sub(r'\s{2,}', ' ', t)
    return t

# ---------- 1. 噪声过滤 ----------
def drop_noise(items, W, H):
    """剔除签名/印章/装订线噪声, 并把页眉页脚分离出来

    返回 (body, header, footer, dropped); dropped 仅供调试图标注用。
    """
    body, header, footer, dropped = [], [], [], []
    for it in items:
        t = clean(it['t'])
        if not t:
            continue
        ry0, ry1 = it['y0'] / H, it['y1'] / H
        rx0, rx1 = it['x0'] / W, it['x1'] / W
        # 手写签名/印章: 边缘区 + 低置信度 + 极短
        if G('drop_stamp') and (ry1 > 0.90 or ry0 < 0.05) and it['s'] < 0.88 and len(t) <= 3:
            dropped.append(dict(it, t=t)); continue
        # 左右页边被裁掉的竖排印记碎片("月""入""R公"...): 正文排不到那儿去。
        # 这里不看置信度 —— 实测这类碎片能到 0.99, 但它们会跨在两行之间, 把
        # 上下两行的包络搭起来并成一行, 危害比读出来的那一两个字大得多。
        if G('drop_stamp') and (rx0 > 0.95 or rx1 < 0.05) and len(t) <= 3:
            dropped.append(dict(it, t=t)); continue
        if G('drop_stamp') and it['s'] < G('stamp_conf') and len(t) <= 4 \
                and not re.match(r'^[\dA-Za-z][\.、)）]?$', t):
            dropped.append(dict(it, t=t))      # 骑缝章/手写签名: 短串+低置信, 印刷体正文均 >0.95
            continue
        rec = dict(it, t=t, rx0=it['x0'] / W, rx1=it['x1'] / W, ry0=ry0, ry1=ry1)
        if G('drop_header') and ry1 < G('header_y') and HEADER_PAT.match(t):
            header.append(rec)
        elif G('drop_footer') and (ry0 > G('footer_y')
                                   or (ry0 > 0.85 and FOOTER_PAT.match(t))):
            footer.append(rec)
        else:
            body.append(rec)
    return body, header, footer, dropped

# ---------- 2. 视觉行聚类 ----------
FRAC_SEP = ''      # 私有区字符: 分式占位, 形如 SEP 分子 SEP 分母 SEP

def _stacked(a, b):
    """a、b 是不是上下叠排的一对(横向叠过半 + 一个明显在另一个上方)"""
    ov = min(a['rx1'], b['rx1']) - max(a['rx0'], b['rx0'])
    if ov <= 0.5 * min(a['rx1'] - a['rx0'], b['rx1'] - b['rx0']):
        return None
    up, lo = (a, b) if a['y0'] < b['y0'] else (b, a)
    if lo['y0'] < up['y0'] + 0.5 * (up['y1'] - up['y0']):
        return None
    return up, lo

def _has_bar(arr, up, lo):
    """两个 item 之间有没有一条分数线: 又细又孤立的横向墨迹

    光靠"这一行墨迹多"不行 —— 密排汉字行本身就能占 58%, 表格框线更是满格。
    分数线的签名是: 峰值覆盖 >60%、连续高覆盖不超过 5px、上下 3-4px 处几乎
    全白。实测 107 页里 4 个真分式全中, 表格框线与汉字行 0 误报。
    """
    H, W = arr.shape
    x0, x1 = int(min(up['x0'], lo['x0'])), int(max(up['x1'], lo['x1']))
    y0 = max(0, int(up['y1'] - 0.35 * (up['y1'] - up['y0'])))
    y1 = min(H, int(lo['y0'] + 0.35 * (lo['y1'] - lo['y0'])))
    w = x1 - x0
    if w < 10 or y1 - y0 < 3:
        return False
    cov = (arr[y0:y1, x0:x1] < 160).sum(1) / w
    k = int(cov.argmax())
    if cov[k] <= 0.6 or sum(1 for c in cov if c > 0.6) > 5:
        return False
    halo = [cov[i] for i in (k - 4, k - 3, k + 3, k + 4) if 0 <= i < len(cov)]
    return max(halo, default=0.0) < 0.3

def find_fracs(lines, png, log=None):
    """把"上下叠排 + 中间一条分数线"的两个 item 合成一个分式 item

    合成后它就是普通的一格, 后面的列判定/续行合并都不用管分式这回事;
    真正的上下叠排留到出 docx 时由 OMML 还原。
    """
    cand = [(ln, p) for ln in lines
            for p in (_stacked(a, b)
                      for n, a in enumerate(ln['items']) for b in ln['items'][n + 1:])
            if p]
    if not cand:
        return lines
    import numpy as np
    from PIL import Image
    arr = np.array(Image.open(png).convert('L'))
    n = 0
    for ln, (up, lo) in cand:
        if up not in ln['items'] or lo not in ln['items']:
            continue                       # 已被前一个分式吃掉
        if not _has_bar(arr, up, lo):
            continue
        merged = dict(up, t=f'{FRAC_SEP}{up["t"]}{FRAC_SEP}{lo["t"]}{FRAC_SEP}',
                      x0=min(up['x0'], lo['x0']), x1=max(up['x1'], lo['x1']),
                      y0=up['y0'], y1=lo['y1'],
                      rx0=min(up['rx0'], lo['rx0']), rx1=max(up['rx1'], lo['rx1']),
                      ry0=up['ry0'], ry1=lo['ry1'], s=min(up['s'], lo['s']))
        ln['items'] = [i for i in ln['items'] if i is not up and i is not lo] + [merged]
        ln['items'].sort(key=lambda d: d['x0'])
        _line_meta(ln)
        n += 1
        _log(log, f'  分式: {up["t"]} / {lo["t"]}')
    if n:
        _log(log, f'  识别到 {n} 个分式, 将以 Word 公式(OMML)输出')
    return lines

def _line_meta(ln):
    """行内 item 定了之后, 算出排版要用的几个派生量"""
    ln['rx0'] = ln['items'][0]['rx0']
    ln['rx1'] = max(i['rx1'] for i in ln['items'])
    ln['ry0'] = ln['items'][0]['ry0']
    ln['h'] = ln['y1'] - ln['y0']
    # 内容起点: 首 item 若是窄编号(如 "1." "1.1"), 基准取下一个 item
    first = ln['items'][0]
    ln['cx0'] = (ln['items'][1]['rx0']
                 if len(ln['items']) > 1 and (first['rx1'] - first['rx0']) < 0.08
                 else first['rx0'])

def group_lines(items):
    """y 区间重叠 > G('line_tol') 的 item 合并为一个视觉行

    行包络会随并入的 item 往下长, 这是有意的: 双行表头里"产品名称"这种单行
    标签在格内垂直居中, 中心跟下面那行对齐, 只有让包络吃下整个表头才并得对。
    代价是包络可能被一个跨两行的噪声 item 搭桥、一路长下去 —— 那条路由
    drop_noise 剔页边碎片堵住, 不在这里收紧。
    """
    items = sorted(items, key=lambda d: (d['y0'], d['x0']))
    lines = []
    for it in items:
        placed = False
        for ln in lines:
            ov = min(ln['y1'], it['y1']) - max(ln['y0'], it['y0'])
            h = min(ln['y1'] - ln['y0'], it['y1'] - it['y0'])
            if h > 0 and ov / h > G('line_tol'):
                ln['items'].append(it)
                ln['y0'] = min(ln['y0'], it['y0']); ln['y1'] = max(ln['y1'], it['y1'])
                placed = True
                break
        if not placed:
            lines.append({'y0': it['y0'], 'y1': it['y1'], 'items': [it]})
    for ln in lines:
        ln['items'].sort(key=lambda d: d['x0'])
        _line_meta(ln)
    return sorted(lines, key=lambda l: l['y0'])

# ---------- 2.5 框线表格: 照框线还原行列, 顺带还原合并单元格 ----------
RULE_DARK = 185     # 框线判深浅: 正文用 160, 框线要放宽 —— Excel 导出的浅灰
                    # 边框实测才 170, 而最深的底纹(蓝色表头)是 195, 卡在中间

def _rules(bw, horiz, min_len, max_thick):
    """形态学开运算取出一个方向上的框线, 返回线段

    核就是"这个方向上连续这么长的一串墨迹"。汉字的横竖笔画再粗也短, 一开
    就没了; 留下来的才是框线。比"整行墨迹占比"稳得多 —— 后者在密排汉字行
    上能到 58%, 跟框线分不开。再卡一道厚度: 底纹色块也能开出长条, 但它厚。

    厚度按"面积 ÷ 长度"算, 不能用外接框 —— 扫描件总带零点几度歪斜, 一条
    1500 px 的横线外接框就有七八像素高, 按外接框判会把整张表的框线全扔掉。
    """
    import cv2
    import numpy as np
    k = np.ones((1, min_len) if horiz else (min_len, 1), np.uint8)
    n, _, stats, _ = cv2.connectedComponentsWithStats(
        cv2.morphologyEx(bw, cv2.MORPH_OPEN, k), 8)
    out = []
    for i in range(1, n):
        x, y, w, h, area = stats[i]
        span = w if horiz else h
        if span < min_len or area / span > max_thick:
            continue
        if horiz:
            out.append({'x0': float(x), 'x1': float(x + w), 'y': y + h / 2})
        else:
            out.append({'y0': float(y), 'y1': float(y + h), 'x': x + w / 2})
    return out

def _anchored(vs, ys, tol):
    """只留两端都搭在横线上的竖线

    汉字的竖笔跟框线一样细一样直, 长度也能碰到阈值; 区别是笔画悬在格子中间,
    两头不接横线。不加这条, 一页正文能"检出"三百多条竖线。
    """
    return [v for v in vs
            if any(abs(v['y0'] - y) <= tol for y in ys)
            and any(abs(v['y1'] - y) <= tol for y in ys)]

def _off_text(segs, items, horiz):
    """再滤一道: 压在文字框里的不是框线, 是笔画

    大号字("基本信息"这种标题)的竖笔能有半格高, 两头离上下框线都很近, 光靠
    锚定滤不掉。但框线画在格与格之间, 不会落进 OCR 的文字框里 —— 用这条分。
    """
    out = []
    for s in segs:
        a, b = (s['x0'], s['x1']) if horiz else (s['y0'], s['y1'])
        c = s['y'] if horiz else s['x']
        n = max(1.0, b - a)
        cover = 0.0
        for it in items:
            lo, hi = (it['x0'], it['x1']) if horiz else (it['y0'], it['y1'])
            p, q = (it['y0'], it['y1']) if horiz else (it['x0'], it['x1'])
            if p - 2 <= c <= q + 2:
                cover = max(cover, (min(b, hi) - max(a, lo)) / n)
        # 看"单个"文字框吃掉多少, 不是累加: 框线常贴着相邻格文字框的边走,
        # 累加会把两侧的框各算一半, 把真框线也判成笔画
        if cover <= 0.8:
            out.append(s)
    return out

def _cluster(vals, tol):
    """一维聚类: 挨得比 tol 近的并成一档, 返回各档中心"""
    out = []
    for v in sorted(vals):
        if out and v - out[-1][-1] <= tol:
            out[-1].append(v)
        else:
            out.append([v])
    return [sum(g) / len(g) for g in out]

def _band(edges, v):
    """v 落在 edges 划出的第几格"""
    k = 0
    for i in range(len(edges) - 1):
        if v >= edges[i]:
            k = i
    return k

def _grid_cells(xs, ys, hs, vs, tol):
    """网格里哪些格是连在一起的: 缺了内部框线就是原件里的合并单元格"""
    nr, nc = len(ys) - 1, len(xs) - 1
    def has_v(r, c):                       # (r,c) 右边有没有竖线
        y0, y1 = ys[r], ys[r + 1]
        return any(abs(v['x'] - xs[c + 1]) <= tol
                   and min(v['y1'], y1) - max(v['y0'], y0) >= 0.6 * (y1 - y0)
                   for v in vs)
    def has_h(r, c):                       # (r,c) 下边有没有横线
        x0, x1 = xs[c], xs[c + 1]
        return any(abs(h['y'] - ys[r + 1]) <= tol
                   and min(h['x1'], x1) - max(h['x0'], x0) >= 0.6 * (x1 - x0)
                   for h in hs)
    used = [[False] * nc for _ in range(nr)]
    cells, merged = [], 0
    for r in range(nr):
        for c in range(nc):
            if used[r][c]:
                continue
            w = 1
            while c + w < nc and not used[r][c + w] and not has_v(r, c + w - 1):
                w += 1
            h = 1
            while (r + h < nr and not used[r + h][c]
                   and all(not has_h(r + h - 1, c + k) for k in range(w))):
                h += 1
            for i in range(r, r + h):
                for j in range(c, c + w):
                    used[i][j] = True
            cells.append([r, c, h, w, ''])
            merged += h > 1 or w > 1
    return {'xs': xs, 'ys': ys, 'cells': cells, 'merged': merged,
            'x0': xs[0], 'y0': ys[0], 'x1': xs[-1], 'y1': ys[-1]}

def find_grids(png, items=(), log=None):
    """页图里所有画了框线的表格; 没有框线返回空, 走原来那套几何判定"""
    import numpy as np
    from PIL import Image
    arr = np.array(Image.open(png).convert('L'))
    H, W = arr.shape
    bw = (arr < RULE_DARK).astype(np.uint8)
    tol = max(5, W // 300)
    thick = max(4, W // 500)
    hs = _off_text(_rules(bw, True, max(40, W // 40), thick), items, True)
    if len(hs) < 2:
        return []
    vs = _anchored(_off_text(_rules(bw, False, max(10, H // 120), thick), items, False),
                   _cluster([h['y'] for h in hs], tol), tol * 2)
    if len(vs) < 2:
        return []
    # 相交的线段属于同一张表: 并查集分组, 一页上多张表各归各的
    par = list(range(len(hs) + len(vs)))
    def find(a):
        while par[a] != a:
            par[a] = par[par[a]]
            a = par[a]
        return a
    for i, h in enumerate(hs):
        for j, v in enumerate(vs):
            if (h['x0'] - tol <= v['x'] <= h['x1'] + tol
                    and v['y0'] - tol <= h['y'] <= v['y1'] + tol):
                ra, rb = find(i), find(len(hs) + j)
                if ra != rb:
                    par[ra] = rb
    groups = collections.defaultdict(lambda: ([], []))
    for i, h in enumerate(hs):
        groups[find(i)][0].append(h)
    for j, v in enumerate(vs):
        groups[find(len(hs) + j)][1].append(v)
    grids = []
    for gh, gv in groups.values():
        if len(gh) < 2 or len(gv) < 2:
            continue
        ys, xs = _cluster([h['y'] for h in gh], tol), _cluster([v['x'] for v in gv], tol)
        # 太窄太扁的多半是签名栏下划线之类, 不当表格
        if len(ys) < 2 or len(xs) < 2 or xs[-1] - xs[0] < 0.15 * W:
            continue
        g = _grid_cells(xs, ys, gh, gv, tol)
        # 只有一个格子的不是表格, 是加了边框的文本框(整页外框也长这样) ——
        # 做成 1x1 表格只会把整段文字塞进一格, 还丢了分段
        if len(g['cells']) < 2:
            continue
        g['rxs'] = [x / W for x in xs]        # 列宽/跨页续表判定都用相对坐标
        grids.append(g)
    grids.sort(key=lambda g: g['y0'])
    if grids:
        _log(log, '  框线表格 ' + ', '.join(
            f'{len(g["ys"])-1}行x{len(g["xs"])-1}列'
            + (f'(合并 {g["merged"]} 处)' if g['merged'] else '') for g in grids))
    return grids

def _join_cell(items, x0, x1):
    """一格里的文字接成一段: 上一行排满了就接着写, 没排满的换行

    格内换行不能一律当续行 —— "手动下单/自动下单"是两条并列的值, 接成
    一行就读不出是两项了。
    """
    parts, prev_full = [], False
    for ln in group_lines(items):
        t = clean(' '.join(i['t'] for i in ln['items']))
        if not t:
            continue
        if parts and prev_full:
            parts[-1] = clean(parts[-1] + ('' if is_zh(parts[-1]) else ' ') + t)
        else:
            parts.append(t)
        prev_full = max(i['x1'] for i in ln['items']) - x0 >= 0.85 * (x1 - x0)
    return '\n'.join(parts)

def fill_grid(g, items):
    """把落在网格里的 item 按中心点归进各个(可能合并的)单元格"""
    owner = {}
    for cell in g['cells']:
        r, c, h, w, _ = cell
        for i in range(r, r + h):
            for j in range(c, c + w):
                owner[(i, j)] = cell
    bag = collections.defaultdict(list)
    for it in items:
        r = _band(g['ys'], (it['y0'] + it['y1']) / 2)
        c = _band(g['xs'], (it['x0'] + it['x1']) / 2)
        cell = owner.get((r, c))
        if cell is not None:
            bag[id(cell)].append(it)
    for cell in g['cells']:
        its = bag.get(id(cell))
        if its:
            cell[4] = _join_cell(its, g['xs'][cell[1]], g['xs'][cell[1] + cell[3]])
    return g

def in_grid(it, grids):
    cx, cy = (it['x0'] + it['x1']) / 2, (it['y0'] + it['y1']) / 2
    return any(g['x0'] <= cx <= g['x1'] and g['y0'] <= cy <= g['y1'] for g in grids)

# ---------- 3. 块划分: 表格区 vs 正文区 ----------
def _max_gap(ln):
    """一行内部最大的横向空白: 表格行的列间距远大于正文里的词间距"""
    its = ln['items']
    return max((b['rx0'] - a['rx1'] for a, b in zip(its, its[1:])), default=0.0)

def _spans(ln):
    """行内 item 合成横向区间: 横向叠掉一半以上的两个 item 是被并成一行的
    上下两行(表头"不含税单价"压着"(元)"那种), 属于同一格, 不能各算一列。

    只是首尾相接不算 —— 第 60 页表头 序号 收在 0.173、数量 正好起于 0.173,
    并掉的话整个数量列就没了。
    """
    out = []
    for it in sorted(ln['items'], key=lambda i: i['rx0']):
        a, b = it['rx0'], it['rx1']
        ov = min(b, out[-1][1]) - a if out else 0.0
        if out and ov > 0 and ov >= 0.5 * min(b - a, out[-1][1] - out[-1][0]):
            out[-1][1] = max(out[-1][1], b)
        else:
            out.append([a, b])
    return out

def _col_starts(lines, tol=0.022):
    """列起点 = 各格左边界的聚类中心, 行数够时要求至少两行对齐才算一列

    不用"找贯穿空白通道"那套: 扫描件里相邻单元格常常几乎贴着
    (实测第 60 页 描述列 收在 0.775, 备注列 起于 0.790, 只差 1.5%),
    通道法会把这两列并掉; 左边界对齐则稳得多。
    """
    sp = sorted((a, b, k) for k, ln in enumerate(lines) for a, b in _spans(ln))
    if not sp:
        return []
    # 格内折行往往比首行缩进几个字, 左边界会另立一列。它整个被同列另一格包着,
    # 就归到那一格的起点上去。"需方：XX ␣␣ 签订时间：YY"这种另一套版面的行跟
    # 表格行只是错开、互不包含, 归不掉 —— 后面由 _crosses 把它分出表格区。
    xs = []
    for a, b, k in sp:
        host = [A for A, B, K in sp if K != k and A <= a - tol and B >= b - 0.01]
        xs.append(min(host) if host else a)
    xs.sort()
    need = 2 if len(lines) > 1 else 1
    groups, cur = [], [xs[0]]
    for x in xs[1:]:
        if x - cur[-1] <= tol:
            cur.append(x)
        else:
            groups.append(cur); cur = [x]
    groups.append(cur)
    starts = [sum(g) / len(g) for g in groups if len(g) >= need]
    # 最左一列常常只有一行有内容(表头缩进、序号列多半空着), 聚类会把它丢掉;
    # 丢了的话 _col_of 会把左边的字全塞进第二列, 所以补回来
    if starts and xs[0] < starts[0] - tol:
        starts.insert(0, xs[0])
    return starts

def _crosses(ln, starts):
    """行内有没有文字横跨列边界 —— 真表格的文字不越格, 整幅宽的散文才越"""
    return any(a < s - 0.02 and b > s + 0.02
               for a, b in _spans(ln) for s in starts)

def _fit(rows):
    """这组行能不能共用一套列: 能则给出列起点, 不能返回 None"""
    starts = _col_starts(rows)
    if len(starts) < 2:
        return None
    return None if any(_crosses(ln, starts) for ln in rows) else starts

def _col_of(x, starts):
    """item 落在哪一列: 取不超过它左边界的最后一个列起点"""
    k = 0
    for i, s in enumerate(starts):
        if x >= s - 0.02:
            k = i
    return k

def _grow(lines, seeds, i):
    """从第 i 行(一个 seed)往下扩表格区, 返回 (末行下标, 列起点, seed 行数)

    往下并一个 seed 之前先试算: 加进来之后整组还得共用同一套列。列位对不上
    的 seed 就是另一回事了, 到此为止 —— 合同抬头"需方：XX ␣␣ 签订时间：YY"
    跟下面的产品明细表列位完全不同, 靠这条才不会被并成一张列数虚高、
    大半格是空的表。夹在中间的单列行是格内续行, 一并收下; 越格的(整幅宽的
    散文)截断。
    """
    n = len(lines)
    cols = _col_starts([lines[i]])
    end, cur, j = i, [lines[i]], i + 1
    while j < n:
        if lines[j]['ry0'] - lines[j - 1]['ry0'] >= 0.05:
            break                                # 行距明显拉开 -> 不是同一张表
        if seeds[j]:
            trial = _fit(cur + [lines[j]])
            if trial is None:
                break
            cols, cur, end = trial, cur + [lines[j]], j
        elif cols and _crosses(lines[j], cols):
            break
        j += 1
    return end, (cols if len(cols) >= 2 else None), len(cur)

def split_blocks(lines):
    """切成 [('tbl', lines, 列起点), ('txt', lines, None), ...]

    表格行(seed)的判据是"行内有超过 G('gutter') 的横向空白"; 表格区怎么长
    见 _grow。表格尾部之后的正文行不收 —— 所以按"最后一个 seed"截断。
    """
    g = G('gutter')
    seeds = [_max_gap(ln) >= g for ln in lines]
    n, blocks, i = len(lines), [], 0
    while i < n:
        if not seeds[i]:
            j = i
            while j < n and not seeds[j]:
                j += 1
            blocks.append(('txt', lines[i:j], None))
            i = j
            continue
        end, cols, rows = _grow(lines, seeds, i)
        seg = lines[i:end + 1]
        if G('tables') and cols and rows >= G('min_tbl_rows'):
            blocks.append(('tbl', seg, cols))
        else:
            blocks.append(('txt', seg, None))
        i = end + 1
    return blocks

def _weave(lines, grids):
    """按 y 在框线表处把正文行切开, 再各自分块 -> 网格与正文保持原来的先后

    不能先整页分块再把网格插进去: 表格里的字被网格拿走后, 表格上下的正文
    变成连续的一整块, 网格无论排在它前面还是后面都是错的。
    """
    if not grids:
        return split_blocks(lines)
    out, cur, gi = [], [], 0
    for ln in lines:                       # group_lines 已按 y 排好
        while gi < len(grids) and ln['y0'] > grids[gi]['y1']:
            out += split_blocks(cur) + [('grid', grids[gi], None)]
            cur, gi = [], gi + 1
        cur.append(ln)
    out += split_blocks(cur) + [('grid', g, None) for g in grids[gi:]]
    return out

def grid_rows(g):
    """网格 -> 文本矩阵(合并块的字写在左上格), 供表头判定/调试用"""
    nr, nc = len(g['ys']) - 1, len(g['xs']) - 1
    m = [[''] * nc for _ in range(nr)]
    for r, c, _h, _w, t in g['cells']:
        m[r][c] = t
    return m

# ---------- 4. 单列块: 合并续行 ----------
def merge_paras(lines):
    """把被 OCR 拆散的续行合并回段落; 未排满的行绝不吸收下一行"""
    # 本页行距中位数 -> 空行阈值(A卷 ~1.5%, B卷 ~3.2%, 写死阈值必误伤其一)
    gaps = sorted(b['ry0'] - a['ry0'] for a, b in zip(lines, lines[1:])
                  if 0 < b['ry0'] - a['ry0'] < 0.08)
    lead = gaps[len(gaps) // 2] if gaps else 0.02
    gap_max = lead * 1.75
    paras = []
    for ln in lines:
        text = clean(' '.join(i['t'] for i in ln['items']))
        if not text:
            continue
        w = ln['rx1'] - ln['rx0']
        cur = {'text': text, 'rx0': ln['rx0'], 'cx0': ln['cx0'],
               'rx1': ln['rx1'], 'h': ln['h'], 'ry0': ln['ry0'],
               'center': abs((ln['rx0'] + ln['rx1']) / 2 - 0.5) < 0.07 and w < 0.5 and ln['rx0'] > 0.25}
        if not paras:
            paras.append(cur); continue
        prev = paras[-1]
        new_block = (
            not (prev['rx1'] > G('full_line')) or              # ★上一行没排满 -> 它已结束
            NUM_START.match(text) or                      # 新编号/新项目符号
            END_PUNCT.search(prev['text']) or             # 上一段已收尾
            is_zh(text) != is_zh(prev['text']) or         # 中英切换 -> 对照的另一半
            abs(cur['cx0'] - prev['cx0']) > G('x_tol') or      # 缩进层级变了
            (cur['ry0'] - prev['ry0']) > gap_max          # 行距明显变大 = 空行
        )
        if new_block:
            paras.append(cur)
        else:
            sep = '' if is_zh(prev['text']) else ' '
            prev['text'] = clean(prev['text'] + sep + text)
            prev['rx1'] = cur['rx1']
            prev['ry0'] = cur['ry0']          # ★ 推进基准行, 否则下一续行会被当空行断开
    return paras

def mark_bullets(paras):
    """OCR 会丢弃 • 图形符号: 用"缩进大于块基准"把列表项找回来

    基准按块(而非整页)取众数: 技术规格那种整块统一缩进的"标签—值"清单,
    块内众数就是它自己, 于是不会被误判成列表; 换成页级基准反而会把它们
    全标成 •。
    """
    if not paras:
        return paras
    base = collections.Counter(round(p['cx0'] * 200) / 200 for p in paras).most_common(1)[0][0]
    for p in paras:
        p['bullet'] = ((p['cx0'] - base) > G('bullet_ind')
                       and not NUM_START.match(p['text'])
                       and not p.get('center'))   # 居中大标题左缩进同样很大, 但不是列表项
    return paras

# ---------- 5. 表格区: 按列起点归位, 单列行并回上一行 ----------
def build_rows(lines, starts):
    """每行的 item 按列起点归位; 只落在一列里的行是单元格续行, 并回上一行

    "跨两列以上 = 新行"这条判据很关键: 第 60 页里 "1 | TMS部分,文件 | TMS提供"
    整行没有序号, 但它确实是新的一行; 而 "含视频线、电话线" 只占描述列一格,
    是上一行的续行。按有没有首列内容来判会把前者判错。
    """
    edges = list(starts[1:]) + [1.0]        # 各列的右边界
    rows, ends = [], []                     # ends: 每行各列最后一个 item 的右端
    for ln in lines:
        cells = [[] for _ in starts]
        right = [0.0] * len(starts)
        for it in ln['items']:
            k = _col_of(it['rx0'], starts)
            cells[k].append(it['t'])
            right[k] = max(right[k], it['rx1'])
        txt = [clean(' '.join(c)) for c in cells]
        filled = [k for k, c in enumerate(txt) if c]
        if not filled:
            continue
        # 表头后面紧跟的单列行不并进表头(它多半是上一页某行的续尾)
        if (len(filled) >= 2 or not rows
                or (len(rows) == 1 and header_like(rows[0]))):
            rows.append(txt); ends.append(right)
            continue
        k = filled[0]
        prev, pend = rows[-1], ends[-1]
        if not prev[k]:
            prev[k] = txt[k]
        elif (pend[k] - starts[k]) / max(1e-6, edges[k] - starts[k]) >= 0.85:
            # 上一行把这一列填满了 85% 以上 -> 是折行, 接着写。
            # 用填充率而不是"离右边界差多少": 各列宽度相差很大, 绝对值定不出统一阈值
            prev[k] = clean(prev[k] + ('' if is_zh(prev[k]) else ' ') + txt[k])
        else:
            prev[k] = prev[k] + '\n' + txt[k]      # 没排满 -> 单元格内的另一条
        pend[k] = right[k]
    return rows

def header_like(row):
    """一行是不是"全是短标签"的样子"""
    cells = [c for c in row if c]
    return len(cells) >= 2 and all(len(c) <= 12 and '\n' not in c for c in cells)

def is_header_row(rows):
    """首行全是短标签 -> 表头(加粗 + 跨页重复); 值列很长的"标签—值"清单不算"""
    return (len(rows) >= 3 and header_like(rows[0])
            and not any(END_PUNCT.search(c) for c in rows[0] if c))

def repeat_header(row):
    """标到 <w:tblHeader>: 表格跨页时 Word 会自动重复这一行"""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)

# ---------- 6. docx 输出 ----------
def set_font(run, zh_font=None, en_font=None, size=None, bold=False):
    zh_font = zh_font or G('zh_font'); en_font = en_font or G('en_font')
    size = G('font_size') if size is None else size
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), zh_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)

M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def omml_frac(num, den, size=None):
    """OMML 分式: Word 里显示成真正的上下叠排, 双击能进公式编辑器

    python-docx 没有公式 API, 直接拼 XML。<m:oMath> 可以像 run 一样挂在
    段落下, 所以正文和单元格都能用。
    """
    def esc(t):
        return (t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))

    def run(t):
        # 跟 Word 自己写出来的形状对齐: m:r 里先 w:rPr(字体/字号) 再 m:t,
        # 字体必须是 Cambria Math —— 换别的字体 Word 不认它是公式字形
        sz = round((G('font_size') if size is None else size) * 2)
        return (f'<m:r><w:rPr>'
                f'<w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/>'
                f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/></w:rPr>'
                f'<m:t xml:space="preserve">{esc(t)}</m:t></m:r>')
    return parse_xml(
        f'<m:oMath xmlns:m="{M_NS}" xmlns:w="{nsmap["w"]}">'
        f'<m:f><m:num>{run(num)}</m:num><m:den>{run(den)}</m:den></m:f></m:oMath>')

def write(p, text, size=None, bold=False):
    """把文本写进段落; 碰到分式占位符就插一段 OMML 公式

    split 之后下标 0,3,6... 是普通文本, 紧跟的两段是分子分母。
    """
    parts = text.split(FRAC_SEP)
    for i in range(0, len(parts), 3):
        if parts[i]:
            set_font(p.add_run(parts[i]), size=size, bold=bold)
        if i + 2 < len(parts):
            p._p.append(omml_frac(parts[i + 1], parts[i + 2], size))
    return p

def add_para(doc, text, level=0, bold=False, size=None, style=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    if style is None and level:
        pf.left_indent = Cm(0.74 * level)
    write(p, text, size=size, bold=bold)
    if not p.runs:                      # 整段就是一个公式: 补一个空 run 好定字号
        set_font(p.add_run(''), size=size, bold=bold)
    return p

def borderless(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none'); e.set(qn('w:sz'), '0')
        borders.append(e)
    tblPr.append(borders)

def heading_level(text):
    """按编号形态判定标题层级; 返回 None 表示正文"""
    t = text.strip()
    if re.search(r'[。；;]$', t):      # 完整句子 -> 正文, 不是标题
        return None
    if re.match(r'^[A-Z]\.\s+[A-Z]', t) and len(t) < 70:
        return 1
    if re.match(r'^\d{1,2}[\.、]\s*\S', t) and len(t) < 90 and not re.match(r'^\d{1,2}\.\d', t):
        return 2
    if re.match(r'^[一二三四五六七八九十]+[、．.]', t) and len(t) < 60:
        return 2
    if re.match(r'^\d{1,2}\.\d{1,2}[\s、．.]', t) and len(t) < 90:
        return 3
    if re.match(r'^\d{1,2}\s+[A-Z][a-z]', t) and len(t) < 70:
        return 3
    return None

TBL_W = 16.0        # 可用正文宽度(cm): A4 21 - 左右边距 2.5x2

def usable_w(doc):
    """当前版面的可用正文宽度(cm); 横向页比竖版宽一半, 表格列宽得跟着变"""
    s = doc.sections[0]
    try:
        return s.page_width.cm - s.left_margin.cm - s.right_margin.cm
    except AttributeError:
        return TBL_W

def _continues(prev, blocks, page_no):
    """本页开头是不是上一页那张表的续表: 列数与列位置都对得上"""
    if not (prev and blocks and blocks[0][0] == 'tbl'):
        return False
    starts = blocks[0][2]
    return (prev['page'] == page_no - 1 and len(prev['starts']) == len(starts)
            and max(abs(a - b) for a, b in zip(prev['starts'], starts)) < 0.05)

def mark_row(table, page_label):
    """跨页续表时把「原第 N 页」做成表内整行, 免得为了插标记把表切断"""
    r = table.add_row()
    cell = r.cells[0]
    for c in r.cells[1:]:
        cell = cell.merge(c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'—— 原第 {page_label} 页 ——')
    set_font(run, size=8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def fill_row(cells, txt, starts, bold_first, header=False):
    for k, (cell, t) in enumerate(zip(cells, txt)):
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        if t:
            write(p, t, size=10, bold=header or (bold_first and k == 0))

def add_table(doc, rows, starts, header):
    table = doc.add_table(rows=0, cols=len(starts))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    borderless(table)
    # 列宽按原件里各列起点的间距分配, 最后一列吃掉剩余宽度
    edges = list(starts) + [1.0]
    span = [edges[i + 1] - edges[i] for i in range(len(starts))]
    total = sum(span) or 1
    W = usable_w(doc)
    for i, c in enumerate(table.columns):
        c.width = Cm(max(1.2, W * span[i] / total))
    bold_first = not header and len(starts) == 2      # 标签—值: 左列即标签
    for n, txt in enumerate(rows):
        r = table.add_row()
        fill_row(r.cells, txt, starts, bold_first, header=header and n == 0)
        if header and n == 0:
            repeat_header(r)
    return table

def cell_text(cell, text, size, bold=False):
    """往单元格里写字; 格内换行写成多个段落, 分式仍走 OMML"""
    cell.text = ''
    p = cell.paragraphs[0]
    for k, part in enumerate(text.split('\n')):
        if k:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        if part:
            write(p, part, size=size, bold=bold)

def grid_size(nc):
    return 10 if nc <= 6 else (9 if nc <= 10 else 8)     # 列多了得缩字号

def fill_grid_cells(table, g, r0, head):
    """把网格的格子填进表格的第 r0 行起; 原件合并的在这里也合并"""
    size = grid_size(len(g['xs']) - 1)
    for r, c, h, w, t in g['cells']:
        cell = table.cell(r0 + r, c)
        if h > 1 or w > 1:
            cell = cell.merge(table.cell(r0 + r + h - 1, c + w - 1))
        cell_text(cell, t, size, bold=head and r0 + r == 0)

def add_grid(doc, g):
    """框线表格 -> 带框线的真表格, 合并单元格照原件合并"""
    nr, nc = len(g['ys']) - 1, len(g['xs']) - 1
    table = doc.add_table(rows=nr, cols=nc)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    rxs = g['rxs']
    total = rxs[-1] - rxs[0] or 1
    W = usable_w(doc)
    for i, col in enumerate(table.columns):
        col.width = Cm(max(0.5, W * (rxs[i + 1] - rxs[i]) / total))
    head = is_header_row(grid_rows(g))
    fill_grid_cells(table, g, 0, head)
    if head:
        repeat_header(table.rows[0])
    return table

def append_grid(table, g):
    """跨页续表: 往上一页那张框线表后面接行, 不另起一张"""
    r0 = len(table.rows)
    for _ in range(len(g['ys']) - 1):
        table.add_row()
    fill_grid_cells(table, g, r0, False)
    return table

def _grid_continues(prev, blocks, page_no):
    """本页开头的框线表是不是上一页那张的续表: 列数与列位置都对得上"""
    if not (prev and blocks and blocks[0][0] == 'grid'):
        return False
    rxs = blocks[0][1]['rxs']
    return (prev['page'] == page_no - 1 and len(prev['rxs']) == len(rxs)
            and max(abs(a - b) for a, b in zip(prev['rxs'], rxs)) < 0.02)

def render_page(doc, blocks, page_label, state=None):
    state = {} if state is None else state
    for bi, (kind, lines, starts) in enumerate(blocks):
        if kind == 'grid':
            if bi == 0 and _grid_continues(state.get('grid'), blocks, page_label):
                table = state['grid']['obj']
                if state.pop('marker', None):
                    mark_row(table, page_label)
                append_grid(table, lines)
            else:
                table = add_grid(doc, lines)
                doc.add_paragraph()
            state['grid'] = {'obj': table, 'rxs': lines['rxs'], 'page': page_label}
            state['last'] = 'grid'
            state.pop('tbl', None)      # 框线表不跟无框线表互相续接
            continue
        if kind == 'tbl':
            rows = build_rows(lines, starts)
            if not rows:
                continue
            # 只有一行且首列为空 = 上页续行漂过来的孤立文字, 不是真表格
            if len(rows) == 1 and not rows[0][0]:
                add_para(doc, ' '.join(c for c in rows[0] if c))
                continue
            prev = state.get('tbl')
            if bi == 0 and _continues(prev, blocks, page_label):
                # 跨页续表: 直接往上一页那张表里加行, 不新起一张
                table = prev['obj']
                if state.pop('marker', None):
                    mark_row(table, page_label)
                for txt in rows:
                    fill_row(table.add_row().cells, txt, starts,
                             not prev['header'] and len(starts) == 2)
            else:
                header = is_header_row(rows)
                table = add_table(doc, rows, starts, header)
                doc.add_paragraph()
                prev = {'header': header}
            state['tbl'] = {'obj': table, 'starts': starts,
                            'page': page_label, 'header': prev['header']}
            state['last'] = 'tbl'
            state.pop('grid', None)
            continue
        else:
            state['last'] = 'txt'
            state.pop('grid', None)     # 中间隔了正文就不是同一张表了
            paras = mark_bullets(merge_paras(lines))
            for i, para in enumerate(paras):
                t = para['text']
                if BULLET.match(t):
                    add_para(doc, BULLET.sub('', t), style='List Bullet'); continue
                if para.get('bullet'):
                    add_para(doc, t, style='List Bullet'); continue
                lv = heading_level(t)
                if lv:
                    add_para(doc, t, bold=True, size=12 - lv, level=0); continue
                # 短行 + 无收尾标点 + 下一行是长正文 => 无编号小标题
                nxt = paras[i + 1] if i + 1 < len(paras) else None
                if (para['rx1'] < 0.58 and not END_PUNCT.search(t) and len(t) < 60
                        and not NUM_START.match(t)
                        and nxt and nxt['rx1'] > G('full_line')):
                    add_para(doc, t, bold=True); continue
                if para.get('center'):
                    p = add_para(doc, t, bold=True, size=13)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                add_para(doc, t, level=1 if para['cx0'] > 0.16 else 0)

# ---------- 6.5 调试图: 把版面判定结果画回页图 ----------
DBG_COLORS = {
    'tbl':    (0, 160, 0),        # 表格区 -> 绿
    'grid':   (200, 0, 200),      # 框线表格的格子 -> 洋红
    'txt':    (230, 130, 0),      # 正文区 -> 橙
    'col':    (0, 190, 190),      # 列起点 -> 青
    'body':   (0, 90, 220),       # 正文行 -> 蓝
    'header': (150, 150, 150),    # 页眉/页脚 -> 灰
    'drop':   (220, 30, 30),      # 被当噪声剔除 -> 红
}

def draw_debug(png, body, header, footer, dropped, blocks, out):
    """生成版面调试图: 一眼看出哪行被当噪声扔了、块怎么划的

    没有这张图, 调阈值只能靠裁图 + 打印坐标硬看。
    """
    from PIL import Image, ImageDraw, ImageFont
    im = Image.open(png).convert('RGB')
    d = ImageDraw.Draw(im)
    W, H = im.size
    try:
        font = ImageFont.load_default(size=max(16, W // 60))
    except TypeError:
        font = ImageFont.load_default()      # Pillow < 10.1 不支持 size

    def box(it, color, width=2):
        d.rectangle([it['x0'], it['y0'], it['x1'], it['y1']], outline=color, width=width)

    for it in header + footer:
        box(it, DBG_COLORS['header'])
    for it in dropped:
        box(it, DBG_COLORS['drop'], 3)
        d.line([it['x0'], it['y0'], it['x1'], it['y1']], fill=DBG_COLORS['drop'], width=2)
    for it in body:
        box(it, DBG_COLORS['body'], 1)

    # 块边界 + 类型标注; 表格区额外画出每一列的起点
    for kind, lines, starts in blocks:
        if not lines:
            continue
        if kind == 'grid':
            g = lines
            for r, c, h, w, _t in g['cells']:
                d.rectangle([g['xs'][c], g['ys'][r], g['xs'][c + w], g['ys'][r + h]],
                            outline=DBG_COLORS['grid'],
                            width=4 if h > 1 or w > 1 else 2)
            d.text((max(2, g['x0']), max(2, g['y0'] - font.size - 8)),
                   f'GRID {len(g["ys"])-1}x{len(g["xs"])-1} '
                   f'merged {g["merged"]}', fill=DBG_COLORS['grid'], font=font)
            continue
        x0 = min(l['items'][0]['x0'] for l in lines)
        x1 = max(max(i['x1'] for i in l['items']) for l in lines)
        y0 = min(l['y0'] for l in lines)
        y1 = max(l['y1'] for l in lines)
        c = DBG_COLORS[kind]
        d.rectangle([x0 - 6, y0 - 6, x1 + 6, y1 + 6], outline=c, width=3)
        # 标签一律 ASCII: PIL 内置位图字体没有中文字形, 画出来是空白方块
        label = f'TABLE {len(starts)}col' if kind == 'tbl' else 'TEXT'
        d.text((max(2, x0 - 6), max(2, y0 - font.size - 8)),
               f'{label} x{len(lines)}', fill=c, font=font)
        for s in starts or []:
            d.line([W * s, y0 - 6, W * s, y1 + 6], fill=DBG_COLORS['col'], width=2)

    # 排满线 / 页眉页脚线
    d.line([W * G('full_line'), 0, W * G('full_line'), H], fill=(200, 0, 200), width=1)
    for y in (G('header_y'), G('footer_y')):
        d.line([0, H * y, W, H * y], fill=(150, 150, 150), width=1)

    legend = [('body', 'body'), ('header', 'header/footer'), ('drop', 'DROPPED'),
              ('tbl', 'TABLE region'), ('col', 'column start'), ('txt', 'TEXT region'),
              ('grid', 'GRID cell')]
    y = 4
    for key, label in legend:
        d.text((4, y), f'[] {label}', fill=DBG_COLORS[key], font=font)
        y += font.size + 4
    d.text((4, y), '[] full_line', fill=(200, 0, 200), font=font)

    os.makedirs(os.path.dirname(out), exist_ok=True)
    im.save(out)
    return out

def page_png(json_path):
    return f'{CACHE}/pages/{os.path.basename(json_path)[:-5]}.png'

def _landscape(pages):
    """原件多数页是横版就出横版 docx —— 横版表格塞进竖版 A4 只能缩成蚂蚁"""
    from PIL import Image
    n = 0
    for f in pages:
        try:
            w, h = Image.open(page_png(f)).size
        except Exception:
            continue
        n += w > h
    return n * 2 > len(pages)

def make_doc(tag, title, pages, progress=None, stop=None, errors=None, log=None):
    doc = Document()
    sec = doc.sections[0]
    if _landscape(pages):
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Cm(29.7), Cm(21)
    else:
        sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    for m in ('top_margin', 'bottom_margin'):
        setattr(sec, m, Cm(2.2))
    for m in ('left_margin', 'right_margin'):
        setattr(sec, m, Cm(2.5))
    st = doc.styles['Normal']
    st.font.name = G('en_font'); st.font.size = Pt(G('font_size'))
    st.element.rPr.rFonts.set(qn('w:eastAsia'), G('zh_font'))

    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(h.add_run(title), size=16, bold=True)

    state = {}
    for idx, f in enumerate(pages):
        _check(stop)
        _tick(progress, 'layout', idx + 1, len(pages), f'重建版面 {idx+1}/{len(pages)}')
        try:
            _page_into(doc, f, idx + 1, state, log)
        except Cancelled:
            raise
        except Exception as e:
            # 单页版面重建失败不能丢掉整篇: 留一条醒目占位, 其余页照常输出
            _record(errors, log, '版面重建', idx + 1, e)
            p = add_para(doc, f'[第 {idx+1} 页解析失败, 已跳过: '
                              f'{type(e).__name__}: {e}]', bold=True)
            p.runs[0].font.color.rgb = RGBColor(0xC0, 0x30, 0x30)
            state.clear()          # 断了就别再往上一页的表里接
    return doc

def analyze_page(json_path, page_no, log=None):
    """一页 OCR 结果 -> 版面块; docx 与 xlsx 两条输出路径共用这一步"""
    from PIL import Image
    d = json.load(open(json_path))
    png = page_png(json_path)
    W, H = Image.open(png).size
    body, header, footer, dropped = drop_noise(d['items'], W, H)
    grids = find_grids(png, body, log) if G('grid_tables') else []
    rest = body
    if grids:
        # 框线表格里的字归给格子, 剩下的才走"按空白猜列"那套
        for g in grids:
            fill_grid(g, [it for it in body if in_grid(it, [g])])
        rest = [it for it in body if not in_grid(it, grids)]
    lines = find_fracs(group_lines(rest), png, log)
    blocks = _weave(lines, grids)
    if G('debug'):
        out = draw_debug(png, body, header, footer, dropped, blocks,
                         f'{CACHE}/debug/{os.path.basename(png)}')
        _log(log, f'  调试图: {out}')
    tbls = [b for b in blocks if b[0] == 'tbl']
    _log(log, f'  第 {page_no} 页: {len(body)} 行正文, {len(blocks)} 个块'
              + (f' (表格 {len(tbls)} 张, 列数 '
                 + '/'.join(str(len(b[2])) for b in tbls) + ')' if tbls else '')
              + (f' (框线表 {len(grids)} 张)' if grids else '')
              + f', 页眉页脚 {len(header)+len(footer)} 行, 噪声剔除 {len(dropped)} 行')
    return blocks

def _page_into(doc, json_path, page_no, state=None, log=None):
    """把一页 OCR 结果重建进 doc; 供 make_doc 逐页 try 包裹"""
    state = {} if state is None else state
    blocks = analyze_page(json_path, page_no, log)
    if G('page_marker'):
        if (_continues(state.get('tbl'), blocks, page_no)
                or _grid_continues(state.get('grid'), blocks, page_no)):
            state['marker'] = True     # 交给续表, 做成表内一行
        else:
            marker = doc.add_paragraph()
            marker.paragraph_format.space_before = Pt(8)
            mr = marker.add_run(f'—— 原第 {page_no} 页 ——')
            set_font(mr, size=8); mr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    render_page(doc, blocks, page_no, state)

# ---------- 6.6 xlsx 输出: 只导表格 ----------
NUM_ONLY = re.compile(r'-?\d{1,10}(\.\d+)?')

def flat_text(t):
    """把分式占位摊平成一行文本: Excel 单元格里没有 OMML 这回事"""
    parts = t.split(FRAC_SEP)
    if len(parts) == 1:
        return t
    wrap = lambda s: s if re.fullmatch(r'[\w.]+', s) else f'({s})'
    out = ''
    for i in range(0, len(parts), 3):
        out += parts[i]
        if i + 2 < len(parts):
            out += f'{wrap(parts[i + 1])}/{wrap(parts[i + 2])}'
    return out

def _cell_val(t):
    """能当数字用的存成数字, 其余存文本; 返回 (值, 数字格式)

    只认 10 位以内、无前导零的整数/小数。合同编号、型号、长串数字一旦被
    Excel 当成数字, 会掉前导零或变成科学计数法, 宁可留成文本。小数位数按
    原件给一个数字格式, 免得 2060000.00 显示成 2060000。
    """
    t = flat_text(t).strip()
    if NUM_ONLY.fullmatch(t) and not re.match(r'-?0\d', t):
        if '.' in t:
            return float(t), '0.' + '0' * len(t.split('.')[1])
        return int(t), None
    return t, None

def _disp_w(t):
    """列宽用的显示宽度: 一个汉字占两格"""
    return sum(2 if ord(c) > 0x2E80 else 1 for c in str(t))

def _mark(a, b):
    return f'—— 原第 {a} 页 ——' if a == b else f'—— 原第 {a}–{b} 页 ——'

def _sheet_grid(ws, g, page_no, st, cont):
    """框线表格写进工作表: 原件合并的格子在这里也合并

    只合并"有字"的块 —— 空白区照原样合并只会留下一大片死区, 反正没边框,
    看不出区别。
    """
    from openpyxl.styles import Alignment, Font, PatternFill
    top = Alignment(vertical='top', wrap_text=True)
    mid = Alignment(horizontal='center', vertical='center')
    nr, nc = len(g['ys']) - 1, len(g['xs']) - 1
    if cont:
        st['mark'].value = _mark(st['p0'], page_no)      # 续表只改标记的页码范围
    else:
        if st['row'] > 3:
            st['row'] += 1
        st['mark'] = ws.cell(st['row'], 1, _mark(page_no, page_no))
        st['mark'].font = Font(size=9, color='FF999999')
        st['mark'].alignment = mid
        if nc > 1:
            ws.merge_cells(start_row=st['row'], start_column=1,
                           end_row=st['row'], end_column=nc)
        st['maxc'] = max(st['maxc'], nc)
        st['row'] += 1
        st['p0'] = page_no
        st['n'] += 1
    head = is_header_row(grid_rows(g)) and not cont
    r0 = st['row']
    for r, c, h, w, t in g['cells']:
        if not t:
            continue
        val, nfmt = _cell_val(t)
        cell = ws.cell(r0 + r, c + 1, val)
        if h > 1 or w > 1:
            # 合并块居中: 跨行跨列的多半是标题/分类名; 长句子仍靠左
            cell.alignment = Alignment(wrap_text=True, vertical='center',
                                       horizontal='center' if len(t) <= 12 else 'left')
            ws.merge_cells(start_row=r0 + r, start_column=c + 1,
                           end_row=r0 + r + h - 1, end_column=c + w)
        else:
            cell.alignment = top
            st['w'][c] = max(st['w'][c], _disp_w(t))
        if nfmt:
            cell.number_format = nfmt
        if head and r == 0:
            cell.font = Font(bold=True)
            cell.fill = PatternFill('solid', fgColor='FFEFEFEF')
    st['row'] = r0 + nr
    st['grid'] = {'rxs': g['rxs'], 'page': page_no}
    st.pop('tbl', None)        # 框线表不跟无框线表互相续接

def _sheet_tables(ws, blocks, page_no, st):
    """把一页里的表格块写进工作表; st 存着跨页续表要接的那张表

    续表不另起一张、也不插页码行 —— 一张表在 Excel 里得是连续矩形, 中间夹
    一行标记, 筛选和透视就废了。页码改写在表头上方那行标记里, 变成「原第
    53–56 页」。
    """
    from openpyxl.styles import Alignment, Font, PatternFill
    top = Alignment(vertical='top', wrap_text=True)
    mid = Alignment(horizontal='center', vertical='center')
    for bi, (kind, lines, starts) in enumerate(blocks):
        if kind == 'grid':
            _sheet_grid(ws, lines, page_no, st,
                        bi == 0 and _grid_continues(st.get('grid'), blocks, page_no))
            continue
        if kind != 'tbl':
            continue
        rows = build_rows(lines, starts)
        # 只有一行且首列为空 = 上页续行漂过来的孤立文字, 不是真表格
        if not rows or (len(rows) == 1 and not rows[0][0]):
            continue
        cont = bi == 0 and _continues(st.get('tbl'), blocks, page_no)
        if cont:
            head = st['tbl']['header']
            st['mark'].value = _mark(st['p0'], page_no)
        else:
            if st['row'] > 3:
                st['row'] += 1                       # 表与表之间空一行
            # 标记行跨整张表的宽度合并居中: 它是表外的一条分隔带, 不是数据,
            # 合并了才不会看着像"第一列有个奇怪的值"
            st['mark'] = ws.cell(st['row'], 1, _mark(page_no, page_no))
            st['mark'].font = Font(size=9, color='FF999999')
            st['mark'].alignment = mid
            if len(starts) > 1:
                ws.merge_cells(start_row=st['row'], start_column=1,
                               end_row=st['row'], end_column=len(starts))
            st['maxc'] = max(st['maxc'], len(starts))
            st['row'] += 1
            st['p0'] = page_no
            head = is_header_row(rows)
            st['n'] += 1
        bold_first = not head and len(starts) == 2      # 标签—值: 左列即标签
        for n, txt in enumerate(rows):
            for c, t in enumerate(txt):
                if not t:
                    continue
                val, nfmt = _cell_val(t)
                cell = ws.cell(st['row'], c + 1, val)
                cell.alignment = top
                if nfmt:
                    cell.number_format = nfmt
                if head and n == 0 and not cont:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill('solid', fgColor='FFEFEFEF')
                elif bold_first and c == 0:
                    cell.font = Font(bold=True)
                st['w'][c] = max(st['w'][c], _disp_w(t))
            st['row'] += 1
        st['tbl'] = {'starts': starts, 'page': page_no, 'header': head}
    if blocks and blocks[-1][0] != 'grid':
        st.pop('grid', None)       # 表后面还有别的东西, 下一页就不是续表了

def make_book(tag, title, pages, progress=None, stop=None, errors=None, log=None):
    """逐页把表格导进一个工作簿; 与 make_doc 并列的另一条输出路径"""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font
    from openpyxl.utils import get_column_letter
    wb = Workbook()
    ws = wb.active
    ws.title = '表格'
    ws.cell(1, 1, title).font = Font(size=14, bold=True)
    st = {'row': 3, 'n': 0, 'maxc': 1, 'w': collections.defaultdict(int)}
    for idx, f in enumerate(pages):
        _check(stop)
        _tick(progress, 'layout', idx + 1, len(pages), f'导出表格 {idx+1}/{len(pages)}')
        try:
            _sheet_tables(ws, analyze_page(f, idx + 1, log), idx + 1, st)
        except Cancelled:
            raise
        except Exception as e:
            # 单页失败不能丢掉整篇: 留一条醒目占位, 其余页照常导出
            _record(errors, log, '版面重建', idx + 1, e)
            c = ws.cell(st['row'], 1, f'[第 {idx+1} 页解析失败, 已跳过: '
                                      f'{type(e).__name__}: {e}]')
            c.font = Font(bold=True, color='FFC03030')
            st['row'] += 2
            st.pop('tbl', None)        # 断了就别再往上一页的表里接
            st.pop('grid', None)
    # 列宽按该列最长的一格给, 但封顶 42 —— 整格都开了自动换行, 再宽只会
    # 让一列独占屏幕, 反而看不见右边的列
    for c, w in st['w'].items():
        ws.column_dimensions[get_column_letter(c + 1)].width = min(42, max(8, w + 2))
    if st['maxc'] > 1:                      # 标题跟着最宽的那张表一起居中
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=st['maxc'])
        ws.cell(1, 1).alignment = Alignment(horizontal='center')
    if not st['n']:
        ws.cell(3, 1, '（本文件没有识别到表格。整页版面请用 Word 输出。）')
    _log(log, f'  共导出 {st["n"]} 张表格')
    return wb

def set_cache(base_dir):
    """缓存放在 PDF 所在目录, 便于 --rebuild 与手工清理 (exe 可能在任意位置运行)"""
    global CACHE
    CACHE = os.path.join(base_dir or '.', '.pdfrec_cache')
    for sub in ('pages', 'ocr', 'debug'):
        os.makedirs(os.path.join(CACHE, sub), exist_ok=True)
    return CACHE

FMTS = ('docx', 'xlsx', 'both')     # both = 同一次识别出两份

def _build(fmt):
    return ('docx', 'xlsx') if fmt == 'both' else (fmt,)

def convert(pdf, progress=None, stop=None, out_dir=None, cfg=None,
            errors=None, log=None, fmt='docx'):
    """转换单个 PDF, 返回输出路径列表

    progress(stage, cur, total, msg)  stage in {render, ocr, layout, done}
    stop() -> True 时中断并抛 Cancelled
    errors: 传入 list 时, 单页失败会被收集进去而不是中断整篇
    log(msg): 逐条执行日志的接收者; 不传则打印到 stdout
    fmt: docx(默认) / xlsx(只导表格) / both
    """
    if fmt not in FMTS:
        raise ValueError(f'fmt 只能是 {"/".join(FMTS)}, 收到 {fmt!r}')
    if cfg:
        CFG.update(cfg)
    if not os.path.isfile(pdf):
        raise FileNotFoundError(pdf)
    pdf = os.path.abspath(pdf)
    set_cache(os.path.dirname(pdf))
    name = os.path.splitext(os.path.basename(pdf))[0]
    tag = re.sub(r'[^0-9A-Za-z\u4e00-\u9fff]', '', name)[:16] or 'DOC'
    global _last_tag
    _last_tag = tag
    _log(log, f'[{name}] 渲染中...')
    pngs = render_pdf(pdf, tag, progress, stop, errors, log)
    if not pngs:
        raise RuntimeError('没有任何页面渲染成功, 文件可能已损坏或被加密')
    _log(log, f'[{name}] {len(pngs)} 页, OCR 识别中(首次启动稍慢)...')
    jsons = run_ocr(pngs, tag, progress, stop, errors, log)
    _log(log, f'[{name}] 重建版面...')
    # 出两份时第二遍跑的是同一套版面判定, 每页那条统计不必再刷一遍
    quiet = lambda m: None if m.lstrip().startswith('第 ') else _log(log, m)
    outs, sink, out_log = [], errors, log
    for kind in _build(fmt):
        make = make_doc if kind == 'docx' else make_book
        obj = make(tag, name, jsons, progress, stop, sink, out_log)
        out = os.path.join(out_dir or os.path.dirname(pdf), f'{name}.{kind}')
        os.makedirs(os.path.dirname(out) or '.', exist_ok=True)
        obj.save(out)
        _log(log, f'  ✓ {out}  ({len(jsons)} 页, {os.path.getsize(out)/1024:.0f} KB)')
        outs.append(out)
        sink, out_log = [], quiet      # 页级错误同理, 只收第一遍那份
    if errors:
        _log(log, f'  ! 本文件有 {len(errors)} 处失败(已跳过), 详见错误详情')
    _tick(progress, 'done', len(jsons), len(jsons), outs[-1])
    return outs

def main():
    ap = argparse.ArgumentParser(
        description='扫描版 PDF -> 保版式可编辑 Word / Excel 表格 (全本地, 无需联网)')
    ap.add_argument('pdfs', nargs='*', help='待转换的扫描版 PDF (可拖拽到程序图标上)')
    ap.add_argument('--rebuild', metavar='DIR', nargs='?', const='.',
                    help='跳过渲染/OCR, 仅用该目录下的缓存重建输出(格式见 --to)')
    ap.add_argument('--debug', action='store_true',
                    help='生成版面调试图到 .pdfrec_cache/debug/')
    ap.add_argument('--to', choices=FMTS, default='docx',
                    help='输出格式: docx(默认, 保版式) / xlsx(只导表格) / both')
    a = ap.parse_args()
    if a.debug:
        CFG['debug'] = True

    if a.rebuild:
        set_cache(a.rebuild)
        tags = sorted({os.path.basename(f).rsplit('_', 1)[0]
                       for f in glob.glob(f'{CACHE}/ocr/*.json')})
        for tag in tags:
            js = sorted(glob.glob(f'{CACHE}/ocr/{tag}_*.json'))
            for kind in _build(a.to):
                out = os.path.join(a.rebuild, f'{tag}.{kind}')
                (make_doc if kind == 'docx' else make_book)(tag, tag, js).save(out)
                print(f'  ✓ {out}  <- {len(js)} 页 (缓存重建)')
        return

    pdfs = a.pdfs
    if not pdfs and FROZEN:
        # 双击运行(未拖文件): 给提示而不是一闪而过
        print('=' * 58)
        print('  扫描版 PDF -> 可编辑 Word    (全本地运行, 不联网)')
        print('=' * 58)
        print('\n用法: 把 PDF 文件直接拖到本程序图标上即可。\n')
        try:
            s = input('或在此粘贴 PDF 路径(多个用空格分隔, 直接回车退出): ').strip()
        except EOFError:
            s = ''
        pdfs = [p.strip(' "\'') for p in re.findall(r'"[^"]+"|\S+', s)] if s else []

    if not pdfs:
        ap.print_help()
    else:
        failed, page_errs = [], []
        for p in pdfs:
            errs = []
            try:
                convert(p, errors=errs, fmt=a.to)
            except Exception as e:
                # 整篇失败也不抛栈退出: 打完整 traceback 供排查, 继续下一个文件
                print(f'  ✗ 转换失败 {p}: {type(e).__name__}: {e}')
                traceback.print_exc()
                failed.append(p)
            page_errs += [dict(e, file=os.path.basename(p)) for e in errs]
        if failed or page_errs:
            print('\n--- 错误汇总 ---')
            for p in failed:
                print(f'  整篇失败: {p}')
            for e in page_errs:
                print(f'  {e["file"]} 第 {e["page"]} 页 {e["stage"]}: {e["error"]}')
    if FROZEN:
        try:
            input('\n全部完成。按回车键关闭...')
        except EOFError:
            pass          # 管道/无 stdin 环境(如 CI)下不阻塞

if __name__ == '__main__':
    main()
